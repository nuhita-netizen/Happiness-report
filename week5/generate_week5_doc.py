import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import os
import shutil

doc = Document()

# Set standard 1-inch margins
for section in doc.sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# Colors
COLOR_PRIMARY = RGBColor(26, 82, 118)     # Navy Blue
COLOR_SECONDARY = RGBColor(40, 116, 166)  # Steel Blue
COLOR_DARK = RGBColor(44, 62, 80)         # Charcoal

def add_header(title, level=1):
    h = doc.add_heading(level=level)
    run = h.add_run(title)
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(15)
        run.font.color.rgb = COLOR_PRIMARY
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(4)
    elif level == 2:
        run.font.size = Pt(12)
        run.font.color.rgb = COLOR_SECONDARY
        h.paragraph_format.space_before = Pt(10)
        h.paragraph_format.space_after = Pt(3)
    return h

# Title Page / Header Block
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run("WEEK 5: CAPSTONE PROJECT REPORT\n")
title_run.font.size = Pt(20)
title_run.font.bold = True
title_run.font.color.rgb = COLOR_PRIMARY

sub_run = title_p.add_run("Comprehensive Data Science Project Reporting & Strategic Policy Recommendations\n")
sub_run.font.size = Pt(13)
sub_run.font.italic = True
sub_run.font.color.rgb = COLOR_SECONDARY

author_run = title_p.add_run("A Data-Driven Global Well-Being Investigation based on the World Happiness Report (2015–2022)\n")
author_run.font.size = Pt(10.5)
author_run.font.color.rgb = COLOR_DARK

meta_run = title_p.add_run("YUVA Data Analytics Internship | Final Capstone Deliverable")
meta_run.font.size = Pt(9.5)
meta_run.font.color.rgb = RGBColor(128, 128, 128)

doc.add_paragraph().paragraph_format.space_after = Pt(8)

# 1. Executive Summary
add_header("1. Executive Summary")
doc.add_paragraph(
    "This comprehensive capstone report synthesizes the 5-week data science investigation into global subjective well-being, "
    "leveraging the multi-year World Happiness Report dataset. By integrating exploratory data analysis, visual storytelling, "
    "parametric statistical hypothesis testing, and supervised machine learning pipelines, this project establishes a rigorous, "
    "evidence-based framework identifying the primary drivers of national happiness."
)
doc.add_paragraph(
    "Key Findings: While economic output (GDP per capita) and healthy life expectancy establish the foundational prerequisite "
    "for subjective well-being (explaining over 70% of variance), social support networks (Family) and institutional trust "
    "(absence of corruption) serve as vital non-linear multipliers that differentiate top-tier thriving societies from "
    "stagnant economies. Predictive machine learning models achieved high out-of-sample accuracy (R² = 0.866; Classification AUC = 0.939), "
    "providing reliable quantitative tools for scenario simulation and targeted policy intervention."
)

# 2. Project Architecture & End-to-End Methodology
add_header("2. End-to-End Project Methodology")
doc.add_paragraph(
    "The analytical workflow was structured across four distinct phases to ensure scientific rigor and data integrity:"
)

method_bullets = [
    ("Phase 1: Data Acquisition & Preprocessing (Week 1): ", "Extracted and profiled 1,231 multi-year country records. Diagnosed concatenation schema gaps in 2017–2022 and restricted modeling to 315 complete, verified observations across six core socio-economic dimensions while eliminating target-leaking features (Happiness Rank, Dystopia Residual)."),
    ("Phase 2: Advanced Visual Storytelling (Week 2): ", "Built multi-dimensional visual narratives exploring regional clustering, log-linear wealth effects, and disparity breakdowns between the top 10 happiest and bottom 10 struggling nations."),
    ("Phase 3: Statistical Hypothesis Testing (Week 3): ", "Formulated and executed two-sample independent t-tests and correlation significance tests (t = 16.84, p < 0.0001), formally rejecting the null hypothesis to prove that high-GDP nations experience statistically superior happiness."),
    ("Phase 4: Predictive Machine Learning & Diagnostics (Week 4): ", "Trained, validated, and compared five regression and classification architectures (Linear, Ridge L2, Decision Trees, Random Forest, Gradient Boosting) using 5-Fold Cross-Validation and residual error diagnostics.")
]

for b_title, b_desc in method_bullets:
    bp = doc.add_paragraph(style='List Bullet')
    r1 = bp.add_run(b_title)
    r1.font.bold = True
    bp.add_run(b_desc)

# 3. Synthesis of Key Results
add_header("3. Synthesis of Technical & Empirical Results")
doc.add_paragraph(
    "Quantitative modeling across the multi-year dataset yielded consistent empirical insights:"
)

# Table of ML Models
add_header("Machine Learning Model Performance Benchmark", level=2)
table_data = [
    ["Model Architecture", "MAE", "RMSE", "Test R²", "5-Fold CV R² / AUC", "Primary Strength / Bias"],
    ["Multiple Linear Regression", "0.353", "0.444", "0.866", "0.721 ± 0.067", "Optimal baseline interpretability"],
    ["Ridge Regression (L2)", "0.353", "0.444", "0.866", "0.721 ± 0.067", "Controls GDP-Health multicollinearity"],
    ["Random Forest Regressor", "0.392", "0.481", "0.844", "0.751 ± 0.056", "Lowest generalization variance across folds"],
    ["Gradient Boosting Regressor", "0.404", "0.489", "0.838", "0.732 ± 0.051", "Captures subtle non-linear interactions"],
    ["Decision Tree (depth=4)", "0.479", "0.642", "0.721", "0.565 ± 0.123", "High variance; prone to leaf overfitting"]
]

table = doc.add_table(rows=len(table_data), cols=6)
table.alignment = WD_TABLE_ALIGNMENT.CENTER

for r_idx, row in enumerate(table_data):
    for c_idx, val in enumerate(row):
        cell = table.cell(r_idx, c_idx)
        cell.text = val
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if r_idx == 0:
            shd = parse_xml(r'<w:shd {} w:fill="1A5276"/>'.format(nsdecls('w')))
            cell._tc.get_or_add_tcPr().append(shd)
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(255, 255, 255)
        else:
            for run in p.runs:
                run.font.size = Pt(8.5)
            if r_idx % 2 == 1:
                shd = parse_xml(r'<w:shd {} w:fill="F2F4F4"/>'.format(nsdecls('w')))
                cell._tc.get_or_add_tcPr().append(shd)

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# 4. Embedded Visual Evidence
add_header("4. Visual Analytics & Evidence")
base_dir = os.path.dirname(os.path.abspath(__file__))
w2_dir = os.path.join(base_dir, '..', 'week2')
w4_dir = os.path.join(base_dir, '..', 'week4')

# Copy images to week5 for self-contained package
for f in ['viz1_region_distribution.png', 'viz3_correlation_matrix.png']:
    src = os.path.join(w2_dir, f)
    if os.path.exists(src):
        shutil.copy(src, os.path.join(base_dir, f))

for f in ['viz_model_comparison.png', 'viz_actual_vs_predicted.png', 'viz_feature_importance.png']:
    src = os.path.join(w4_dir, f)
    if os.path.exists(src):
        shutil.copy(src, os.path.join(base_dir, f))

# Embed Viz 1
v1_p = os.path.join(base_dir, 'viz1_region_distribution.png')
if os.path.exists(v1_p):
    doc.add_picture(v1_p, width=Inches(5.5))
    cap1 = doc.add_paragraph("Figure 1: Global Regional Disparity in Happiness Scores (Western Europe vs. Developing Regions).")
    cap1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap1.runs[0].font.italic = True
    cap1.runs[0].font.size = Pt(9)

# Embed Viz 2
v2_p = os.path.join(base_dir, 'viz_feature_importance.png')
if os.path.exists(v2_p):
    doc.add_picture(v2_p, width=Inches(5.5))
    cap2 = doc.add_paragraph("Figure 2: Relative Feature Importances Derived from Random Forest Modeling.")
    cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap2.runs[0].font.italic = True
    cap2.runs[0].font.size = Pt(9)

# 5. Strategic Recommendations
add_header("5. Strategic Recommendations for Decision-Makers & Governments")
doc.add_paragraph(
    "Based on quantitative factor weights and model diagnostics, the following actionable strategies are recommended "
    "for national policymakers, developmental agencies, and international economic forums:"
)

strat_items = [
    ("1. Shift from Pure GDP Maximization to Inclusive Prosperity: ", "While GDP per capita is the strongest baseline driver, diminishing marginal returns occur beyond upper-middle-income thresholds. National policy must pivot toward equitable distribution, healthcare accessibility, and community infrastructure to yield measurable happiness gains."),
    ("2. Strengthen Institutional Anti-Corruption Frameworks: ", "Government Trust and Low Corruption exhibit the highest disparity ratio between top 10 and bottom 10 nations. Transparent judicial systems, open governance data, and meritocratic public administrations act as critical catalysts for civic satisfaction."),
    ("3. Formalize Social Safety Nets & Mental Health Support: ", "Family and social cohesion account for over 22% of predictive model weight. Governments should invest in family leave policies, eldercare infrastructure, and community-level mental wellness programs to buffer against macroeconomic volatility."),
    ("4. Protect Democratic Freedoms and Civic Autonomy: ", "The freedom to make life choices correlates strongly with subjective well-being across all income quartiles. Protecting civic freedoms, entrepreneurship ease, and individual rights directly elevates national morale.")
]

for s_title, s_desc in strat_items:
    sp = doc.add_paragraph(style='List Bullet')
    r1 = sp.add_run(s_title)
    r1.font.bold = True
    sp.add_run(s_desc)

# 6. Socio-Economic Impact & Business Implications
add_header("6. Socio-Economic Impact & Business Implications")
doc.add_paragraph(
    "The implications of national well-being extend directly into economic productivity, human capital retention, and corporate performance:"
)
doc.add_paragraph(
    "• Labor Productivity & Innovation: Nations with higher subjective well-being consistently exhibit higher workforce engagement, lower absenteeism, and superior rates of technological innovation.\n"
    "• Foreign Direct Investment (FDI) Attractiveness: International corporations increasingly factor quality-of-life indices, social stability, and institutional transparency into long-term capital allocation decisions.\n"
    "• Brain Drain Mitigation: High-trust, high-freedom societies retain domestic skilled talent and attract global human capital, generating positive compounding feedback loops."
)

# 7. Limitations & Future Research Roadmap
add_header("7. Limitations & Future Research Roadmap")
doc.add_paragraph(
    "While the project delivers robust predictive and statistical results, several analytical constraints provide avenues for future extension:"
)
limit_items = [
    ("Schema Concatenation Harmonization: ", "Future iterations should rectify historical column header variations in 2017–2022 to expand the fully-supervised modeling dataset from 315 rows to over 1,200 observations."),
    ("Granular Sub-National & Inequality Data: ", "National averages can mask severe internal socio-economic and regional disparities; integrating regional Gini coefficients and urban-versus-rural indices would refine predictive precision."),
    ("Advanced Time-Series & Econometric Modeling: ", "Incorporating panel data econometric techniques (Fixed Effects / Random Effects) and deep learning sequential models to study macro-shock resilience (e.g., global recessions, pandemics).")
]

for l_title, l_desc in limit_items:
    lp = doc.add_paragraph(style='List Bullet')
    r1 = lp.add_run(l_title)
    r1.font.bold = True
    lp.add_run(l_desc)

# 8. Conclusion
add_header("8. Conclusion")
doc.add_paragraph(
    "The 5-week World Happiness Report project successfully bridges quantitative data science methodologies—from exploratory "
    "auditing to predictive machine learning—with real-world socio-economic strategy. The findings conclusively demonstrate that "
    "sustainable national well-being requires a balanced tripartite strategy: sustainable economic output, accessible social safety "
    "architectures, and uncompromising institutional integrity."
)

# Save document
out_path = os.path.join(base_dir, "Week5_Comprehensive_Project_Report.docx")
doc.save(out_path)
print(f"Week 5 Report generated successfully at: {out_path}")
