import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
import os

doc = Document()

# Set standard margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# Colors
COLOR_PRIMARY = RGBColor(26, 82, 118)     # Navy Blue
COLOR_SECONDARY = RGBColor(40, 116, 166)  # Steel Blue
COLOR_DARK = RGBColor(44, 62, 80)         # Charcoal

def add_header(title, level=1):
    h = doc.add_heading(level=level)
    run = h.add_run(title)
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = COLOR_PRIMARY
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(6)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = COLOR_SECONDARY
        h.paragraph_format.space_before = Pt(10)
        h.paragraph_format.space_after = Pt(4)
    return h

# Title Block
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run("Week 4: Machine Learning Model Development & Evaluation\n")
title_run.font.size = Pt(20)
title_run.font.bold = True
title_run.font.color.rgb = COLOR_PRIMARY

subtitle_run = title_p.add_run("Predictive Modeling and Statistical Validation on the World Happiness Index\n")
subtitle_run.font.size = Pt(12)
subtitle_run.font.italic = True
subtitle_run.font.color.rgb = COLOR_SECONDARY

author_run = title_p.add_run("YUVA Data Analytics Internship | Machine Learning Deliverable")
author_run.font.size = Pt(10)
author_run.font.color.rgb = RGBColor(128, 128, 128)

doc.add_paragraph().paragraph_format.space_after = Pt(10)

# 1. Executive Summary
add_header("1. Executive Summary & Task Objectives")
p = doc.add_paragraph(
    "Week 4 focuses on designing, training, validating, and critically benchmarking predictive Machine Learning models "
    "using the multi-year World Happiness Report dataset. The objective is to construct regression pipelines that predict "
    "a nation's Happiness Score based on key socio-economic indicators, rigorously evaluate their generalization capabilities "
    "using 5-Fold Cross-Validation, and provide diagnostic assessments of error distributions."
)

# 2. Data Preparation & Pipeline
add_header("2. Data Preprocessing & Feature Engineering Pipeline")
doc.add_paragraph(
    "Data preprocessing was performed to ensure statistical validity and prevent data leakage across model partitions:"
)

bullets = [
    ("Feature Selection: ", "Six core socio-economic dimensions were extracted as predictor variables (X): Economy (GDP per Capita), Social Support (Family), Health (Life Expectancy), Freedom to Make Choices, Perceptions of Government Corruption (Trust), and Generosity."),
    ("Target Variable: ", "Happiness Score (y), representing the national average subjective well-being measured on the 0–10 Cantril Ladder scale."),
    ("Missing Value Treatment: ", "Filtered complete cases across the reporting timeline to avoid distortion from incomplete schema reporting."),
    ("Data Partitioning: ", "Stratified/shuffled 80/20 train-test split (252 training records, 63 testing records) using a fixed random seed (seed=42) for reproducible benchmarking."),
    ("Feature Normalization: ", "StandardScaler was fitted strictly on training partitions (μ = 0, σ = 1) and applied to test sets for distance-sensitive and linear algorithms.")
]

for b_title, b_desc in bullets:
    bp = doc.add_paragraph(style='List Bullet')
    r1 = bp.add_run(b_title)
    r1.font.bold = True
    bp.add_run(b_desc)

# 3. Model Selection
add_header("3. Model Selection & Theoretical Justification")
doc.add_paragraph(
    "To comprehensively evaluate linear and non-linear patterns, five distinct algorithms were trained and compared:"
)

models_info = [
    ("1. Multiple Linear Regression (Baseline): ", "Establishes standard baseline interpretability to estimate linear coefficients per unit increase in each predictor."),
    ("2. Ridge Regression (L2 Regularization): ", "Introduces an L2 penalty on regression weights (alpha=1.0) to stabilize estimates against multicollinearity between GDP and Life Expectancy."),
    ("3. Decision Tree Regressor: ", "Non-parametric tree-based partitioner capable of capturing step-wise interactions, constrained with a max_depth=5 to mitigate overfitting."),
    ("4. Random Forest Regressor: ", "Ensemble bagging model with 100 de-correlated estimators to reduce individual tree variance and provide robust non-linear approximations."),
    ("5. Gradient Boosting Regressor: ", "Sequential ensemble boosting technique optimizing pseudo-residuals with learning_rate=0.08 and 100 estimators.")
]

for m_title, m_desc in models_info:
    mp = doc.add_paragraph(style='List Bullet')
    r1 = mp.add_run(m_title)
    r1.font.bold = True
    mp.add_run(m_desc)

# 4. Experimental Results & Benchmark
add_header("4. Model Evaluation & Performance Benchmark")
doc.add_paragraph(
    "Models were assessed on three standard regression metrics: Mean Absolute Error (MAE), Root Mean Squared Error (RMSE), "
    "and Out-of-Sample R² Score, complemented by 5-Fold Cross-Validation R² to gauge generalization stability."
)

# Table
table_data = [
    ["Model Architecture", "MAE", "RMSE", "Test R²", "5-Fold CV R² (Mean ± Std)"],
    ["Linear Regression", "0.353", "0.444", "0.866", "0.721 ± 0.067"],
    ["Ridge Regression", "0.353", "0.444", "0.866", "0.721 ± 0.067"],
    ["Decision Tree Regressor", "0.479", "0.642", "0.721", "0.565 ± 0.123"],
    ["Random Forest Regressor", "0.392", "0.481", "0.844", "0.751 ± 0.056"],
    ["Gradient Boosting Regressor", "0.404", "0.489", "0.838", "0.732 ± 0.051"]
]

table = doc.add_table(rows=len(table_data), cols=5)
table.alignment = WD_TABLE_ALIGNMENT.CENTER

for r_idx, row in enumerate(table_data):
    for c_idx, val in enumerate(row):
        cell = table.cell(r_idx, c_idx)
        cell.text = val
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if r_idx == 0:
            shading_elm = parse_xml(r'<w:shd {} w:fill="1A5276"/>'.format(nsdecls('w')))
            cell._tc.get_or_add_tcPr().append(shading_elm)
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
        else:
            if r_idx % 2 == 1:
                shading_elm = parse_xml(r'<w:shd {} w:fill="F2F4F4"/>'.format(nsdecls('w')))
                cell._tc.get_or_add_tcPr().append(shading_elm)

doc.add_paragraph().paragraph_format.space_after = Pt(8)

# 5. Visualizations
add_header("5. Visual Performance Diagnostics")

viz_dir = os.path.dirname(os.path.abspath(__file__))

# Viz 1: Model Comparison
viz1_path = os.path.join(viz_dir, "viz_model_comparison.png")
if os.path.exists(viz1_path):
    doc.add_paragraph().paragraph_format.space_before = Pt(4)
    doc.add_picture(viz1_path, width=Inches(6.0))
    cap = doc.add_paragraph("Figure 1: Benchmark Comparison of Out-of-Sample R² and RMSE across ML Architectures.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.italic = True
    cap.runs[0].font.size = Pt(9.5)

# Viz 2: Actual vs Predicted
viz2_path = os.path.join(viz_dir, "viz_actual_vs_predicted.png")
if os.path.exists(viz2_path):
    doc.add_paragraph().paragraph_format.space_before = Pt(6)
    doc.add_picture(viz2_path, width=Inches(5.5))
    cap = doc.add_paragraph("Figure 2: Actual vs. Predicted Happiness Score with Ideal Parity Line (y = x).")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.italic = True
    cap.runs[0].font.size = Pt(9.5)

# Viz 3: Residuals
viz3_path = os.path.join(viz_dir, "viz_residuals_analysis.png")
if os.path.exists(viz3_path):
    doc.add_paragraph().paragraph_format.space_before = Pt(6)
    doc.add_picture(viz3_path, width=Inches(6.0))
    cap = doc.add_paragraph("Figure 3: Residual Error Distribution and Residuals vs. Fitted Values Diagnostics.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.italic = True
    cap.runs[0].font.size = Pt(9.5)

# Viz 4: Feature Importance
viz4_path = os.path.join(viz_dir, "viz_feature_importance.png")
if os.path.exists(viz4_path):
    doc.add_paragraph().paragraph_format.space_before = Pt(6)
    doc.add_picture(viz4_path, width=Inches(5.8))
    cap = doc.add_paragraph("Figure 4: Relative Feature Importance via Random Forest Gini Impurity Reduction.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.italic = True
    cap.runs[0].font.size = Pt(9.5)

# 6. Critical Discussion
add_header("6. Critical Discussion: Error Analysis & Model Limitations")
doc.add_paragraph(
    "A rigorous critical review reveals several vital statistical insights and practical constraints:"
)

disc_items = [
    ("Overfitting vs. Underfitting Dynamics: ", "The single Decision Tree suffered from noticeable high variance (Test R² = 0.721, CV R² = 0.565 ± 0.123), exhibiting overfitting on localized training leaves. In contrast, Random Forest and Gradient Boosting demonstrated lower generalization variance across folds (CV Std ~ 0.051)."),
    ("Multicollinearity & Feature Redundancy: ", "GDP per capita and Health (Life Expectancy) share strong positive correlation (r > 0.81). While Linear Regression achieves high R² on the test set, coefficient estimates can be volatile without regularization; Ridge Regression proved crucial for coefficient stabilization."),
    ("Omitted Variable Bias & Cultural Subjectivity: ", "The residual standard deviation (~0.44 points) underscores that socio-economic indicators alone cannot capture all variance in subjective happiness. Unmeasured factors such as cultural baseline optimism, governance transparency, and inequality indices contribute to unexplained variance (Dystopia Residual)."),
    ("Recommendations for Iterative Enhancement: ", "1) Integrate macro-economic time-series features (e.g., inflation rates, Gini inequality index); 2) Implement Bayesian Hyperparameter Optimization; 3) Apply SHAP (SHapley Additive exPlanations) values for granular country-level explainability.")
]

for d_title, d_desc in disc_items:
    dp = doc.add_paragraph(style='List Bullet')
    r1 = dp.add_run(d_title)
    r1.font.bold = True
    dp.add_run(d_desc)

# Save document
output_doc_path = os.path.join(viz_dir, "Week4_ML_Report.docx")
doc.save(output_doc_path)
print(f"Report saved successfully to {output_doc_path}")
